# -*- coding: utf-8 -*-
# time: 2023-12-19 21:49
# file: hetong.py
# author: 罗永贵

from os import makedirs
from os.path import exists
import logging
import time
from docx.oxml.ns import qn
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # 设置对象居中、对齐等
from docx.shared import Pt  # 设置像素、缩进等
from docx.shared import RGBColor  # 设置字体颜色
import requests
from bs4 import BeautifulSoup
import re

# 测试耗时
def test_run_time(func):
    def inner(*args, **kwargs):
        t_start = time.time()
        func(*args, **kwargs)
        t_end = time.time()
        print(f'一共花费了{t_end-t_start}秒时间')
        # return res
    return inner

# 配置日志基本设置
logging.basicConfig(
    level=logging.DEBUG,     # 设置为DEBUG级别
    format='%(asctime)s - %(levelname)s - %(message)s',  # 设置日志格式
    handlers=[
        logging.StreamHandler()  # 输出到控制台
    ]
)
# 常量设置
RESULTS_DIR = 'results'
exists(RESULTS_DIR) or makedirs(RESULTS_DIR)
BASE_URL_HT = 'https://hetong.110.com/hetong_%s.html'
PAGE_NUM = 10
HEADERS = {
    'Cookie':'Hm_lvt_d46a84bd6b21a59c09ac3065ce41e52e=1702996840,1702997042,1703065560,1703122926; Hm_lpvt_d46a84bd6b21a59c09ac3065ce41e52e=1703167533; _af_capnum_=10',
    'Referer': 'https://www.baidu.com/s?ie=utf-8&f=8&rsv_bp=1&rsv_idx=1&tn=62095104_17_oem_dg&wd=%E5%90%88%E5%90%8C%E8%8C%83%E6%9C%AC110.com&fenlei=256&oq=%25E5%2590%2588%25E5%2590%258C%25E8%258C%2583%25E6%259C%25AC110.coom&rsv_pq=b48c2a6500213a39&rsv_t=f8e15ROnjbXQ%2BP6IqgqdegeqBHDHwWHCg4ZZHh648qBvHY1eUDuVarCtT8xLExLkjGTlrsZeEQjT&rqlang=cn&rsv_dl=tb&rsv_enter=1&rsv_btype=t&inputT=240&rsv_sug3=30&rsv_sug1=24&rsv_sug7=100&rsv_sug2=0&rsv_sug4=2148',
    'User-Agent':'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
}

# 获取HTML页面
def scrape_page(url):
    logging.debug(f'查询到url:{url}')
    try:
        response = requests.get(url, headers=HEADERS, allow_redirects=False)
        if response.status_code == 200:
            if len(response.text) < 100:
                return  # 跳过当前地址
            else:
                return response.text
        logging.error(f'get onvalid status code {response.status_code} while scraping {url}')
    except requests.RequestException:
        '''exc_info=True: 这是 logging.error 方法的一个参数，它表示在记录日志时同时记录异常信息（即堆栈跟踪）'''
        logging.error(f'error occurred while scraping {url}', exc_info=True)

def scrape_index(page):
    index_url = f'{BASE_URL_HT}/hetong_{page}.html'
    return scrape_page(index_url)

# 解析HTML页面
def parse_index(html):
    if html is not None:
        soup = BeautifulSoup(html, "html.parser")
        title = soup.find_all(name=["h1"], attrs={"id": "articleTitle"})
        content = soup.find_all(name=["div"], attrs={"id": "articleBody"})
        print(title[0].text)
        return {
            'title': title[0].text,
            'content': content[0].text
        }
    else:
        return

# 信息保存
def sava_data(name, content):
    '''
    :param name: 标题
    :param content: 正文内容
    :return:
    '''
    doc = Document()
    # 标题内容获取及格式设置
    title = doc.add_heading(level=1)  # 创建一级标题
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
    title_run = title.add_run(name, )  # 添加标题
    title_run.font.size = Pt(24)  # 设置标题大小
    title_run.font.name = 'Times New Roman'  # 标题英文字体
    title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋简体')  # 标题中文字体
    title_run.font.color.rgb = RGBColor(0, 0, 0)  # 字体颜色
    doc.styles['Normal'].font.name = '宋体'  # 设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')  # 设置中文字体
    '''
    正文内容获取及格式设置
    '''
    paragraph = doc.add_paragraph()  # 添加一个段落
    paragraph_run = paragraph.add_run(content, )
    paragraph.paragraph_format.space_before = Pt(20)  # 段前20磅
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置两端对齐方式
    paragraph.paragraph_format.line_spacing = Pt(28.8)  # 设置行间距
    paragraph_run.font.size = Pt(15)  # 设置字号
    TITLE = re.sub('([^\u4e00-\u9fa5\d])', '', name)  # 删除标题中的特殊字符

    save_name = f'{RESULTS_DIR}/{TITLE}.docx'
    doc.save(save_name)

@test_run_time
def main():
    for num in range(1, PAGE_NUM+1):
        # 获取第一个页面text
        next_url = BASE_URL_HT % num
        html = scrape_page(next_url)
        # html = scrape_index(page)
        # 解析第一个html页面-获取其中的url地址
        parse_this = parse_index(html)
        if parse_this is not None:
            title = parse_this.get('title')
            content = parse_this.get('content')
            logging.debug('saving data to docx file')
            sava_data(name=title, content=content)
            logging.debug('data saved successfully')
        else:
            print("html页面已被删除！！")
        time.sleep(2)

if __name__ == '__main__':
    main()

